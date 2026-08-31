#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""20260831 主稿待落盘两项 (需文件解锁): S3 Panel A 标签 + Figure3/5 重嵌 300dpi."""
import copy
import time

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = r"F:/MIMIC3_1/V13/manuscript/PLF_OGT_讨论重构_图注精简与A4表格优化版_20260831.docx"


def red_rpr(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("• 图 2｜"):
            for r in p.runs:
                if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == "FF0000":
                    return copy.deepcopy(r._r.find(qn("w:rPr")))
    return None


def span_replace_red(p, old, new, rpr):
    full = p.text
    pos = full.find(old)
    if pos < 0:
        return False
    end = pos + len(old)
    off = 0
    first = True
    for r in list(p.runs):
        rl = len(r.text)
        rs, re_ = off, off + rl
        if re_ <= pos or rs >= end:
            off = re_
            continue
        s, e = max(pos - rs, 0), min(end - rs, rl)
        if first:
            head, tail = r.text[:s], r.text[e:]
            r.text = head
            anchor = r._r
            if tail:
                tr = OxmlElement("w:r")
                if anchor.find(qn("w:rPr")) is not None:
                    tr.append(copy.deepcopy(anchor.find(qn("w:rPr"))))
                t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = tail
                tr.append(t)
                anchor.addnext(tr)
            if rpr is not None:
                mid = OxmlElement("w:r"); mid.append(copy.deepcopy(rpr))
                t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = new
                mid.append(t)
                anchor.addnext(mid)
            else:
                r.text = head + new
            first = False
        else:
            r.text = r.text[:s] + r.text[e:]
        off = re_
    return True


def run():
    d = docx.Document(F)
    rpr = red_rpr(d)
    # 1) S3 Panel A 标签
    cell = d.tables[6].rows[1].cells[0]
    ok = any(span_replace_red(p, "全部可评价锚点（n=64,715）",
                              "锚点与终点双时点共同有效锚点（n=64,715）", rpr)
             for p in cell.paragraphs if "全部可评价锚点（n=64,715）" in p.text)
    assert ok, "表6 R1C0 未命中"
    # 2) Figure3/5 重嵌 300dpi
    k = 0
    for p in d.paragraphs:
        for b in p._p.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
            rid = b.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            k += 1
            if k in (3, 5):
                d.part.related_parts[rid]._blob = open(
                    rf"F:/MIMIC3_1/V13/manuscript/投稿图_083113_44/Figure{k}.png", "rb").read()
    d.save(F)
    return k


for i in range(12):
    try:
        n = run()
        print(f"完成: S3 标签 + Figure3/5 重嵌 (共扫 {n} 张图)")
        break
    except PermissionError:
        time.sleep(10)
else:
    print("LOCKED: 文件仍被占用")
