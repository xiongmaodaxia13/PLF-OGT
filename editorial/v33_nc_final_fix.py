#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""NC V3 终版修复 (2026-08-31): 三阻断项 + GitHub 声明时态.

1) Figure3/5 重嵌 300dpi 终版
2) 图5 图题补「部分」
3) R4 (P101) 补 80.1% 循环份额句
4) Code availability: "will be organized into" → "is available in"
修订色 = 本文件惯例蓝 0000FF。文件锁定时轮询重试。
"""
import copy
import time

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = r"F:/MIMIC3_1/V13/manuscript/PLF_OGT_讨论重构_NC署名与声明补充版_V3_20260831.docx"
FIGD = r"F:/MIMIC3_1/V13/manuscript/投稿图_083113_44"

SENT_801 = "其中，循环分量约占六器官 6 h MAE 改善总和的 80.1%（差值基于未舍入原始值计算）。"


def blue_rpr(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == "0000FF":
                return copy.deepcopy(r._r.find(qn("w:rPr")))
    return None


def span_replace(p, old, new, rpr):
    full = p.text
    pos = full.find(old)
    if pos < 0:
        return False
    end = pos + len(old)
    off = 0
    first = True
    for r in list(p.runs):
        rl = len(r.text)
        rs, re_ = off, off + rl
        if re_ <= pos or rs >= end:
            off = re_
            continue
        s, e = max(pos - rs, 0), min(end - rs, rl)
        if first:
            head, tail = r.text[:s], r.text[e:]
            r.text = head
            anchor = r._r
            if tail:
                tr = OxmlElement("w:r")
                if anchor.find(qn("w:rPr")) is not None:
                    tr.append(copy.deepcopy(anchor.find(qn("w:rPr"))))
                t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = tail
                tr.append(t)
                anchor.addnext(tr)
            mid = OxmlElement("w:r")
            if rpr is not None:
                mid.append(copy.deepcopy(rpr))
            else:
                pr = anchor.find(qn("w:rPr"))
                if pr is not None:
                    mid.append(copy.deepcopy(pr))
            t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = new
            mid.append(t)
            anchor.addnext(mid)
            first = False
        else:
            r.text = r.text[:s] + r.text[e:]
        off = re_
    return True


def run():
    d = docx.Document(F)
    rpr = blue_rpr(d)
    log = []
    # 1) 重嵌
    k = 0
    for p in d.paragraphs:
        for b in p._p.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
            rid = b.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            k += 1
            if k in (3, 5):
                d.part.related_parts[rid]._blob = open(rf"{FIGD}/Figure{k}.png", "rb").read()
    log.append(f"1) Figure3/5 重嵌 300dpi (扫描 {k} 张)")
    # 2) 图5 题补「部分」
    for p in d.paragraphs:
        if p.text.strip().startswith("• 图 5｜跨队列独立再开发的方向性重复"):
            span_replace(p, "跨队列独立再开发的方向性重复。", "跨队列独立再开发重现部分方向性结果。", rpr)
            log.append("2) 图5 图题已补「部分」")
            break
    # 3) 80.1% 句 (P101, 插在「总 SOFA MAE 和 AUROC 的结果方向一致」之前)
    ok = False
    for p in d.paragraphs:
        if "总 SOFA MAE 和 AUROC 的结果方向一致" in p.text and "0.357" in p.text:
            ok = span_replace(p, "总 SOFA MAE 和 AUROC 的结果方向一致", SENT_801 + "总 SOFA MAE 和 AUROC 的结果方向一致", rpr)
            break
    log.append(f"3) 80.1% 句: {'已插入' if ok else '!!未命中'}")
    # 4) GitHub 时态
    ok4 = False
    for p in d.paragraphs:
        if "will be organized into a reproducible research repository" in p.text:
            ok4 = span_replace(p, "will be organized into a reproducible research repository at",
                               "is available in a reproducible research repository at", rpr)
            break
    log.append(f"4) Code availability 时态: {'已改' if ok4 else '!!未命中'}")
    assert all("!!" not in x for x in log), log
    d.save(F)
    return log


for i in range(12):
    try:
        for line in run():
            print(line)
        print("已保存")
        break
    except PermissionError:
        time.sleep(10)
else:
    print("LOCKED")
