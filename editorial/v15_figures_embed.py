#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""V15 图片嵌入手术: 删除旧图 2/3 图片, 在图 2–5 图注前嵌入新图 (PNG 300dpi, 180mm 宽).

终态: 文档末尾图注区块 = 图1(原图+原注) + 图2-5(新图在上、注在下, 与图 1 布局一致)。
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.shared import Cm
from docx.oxml.ns import qn

BASE = Path(r"F:/MIMIC3_1/V13/manuscript")
F = BASE / "PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260830.docx"
FIGD = BASE / "figures"


def main():
    doc = docx.Document(str(F))

    # 1) 定位文档末尾区块的图片段落与图注段落
    cap_idx = {}   # 图号 -> 段落对象 (末尾区块, 索引>185)
    img_paras = []  # 含 w:drawing 的段落
    for i, p in enumerate(doc.paragraphs):
        if i <= 185:
            continue
        t = p.text.strip()
        if t.startswith("图 1｜"):
            cap_idx[1] = (i, p)
        elif t.startswith("图 2｜总 SOFA"):
            cap_idx[2] = (i, p)
        elif t.startswith("图 3｜锚点后"):
            cap_idx[3] = (i, p)
        elif t.startswith("图 4｜概念"):
            cap_idx[4] = (i, p)
        elif t.startswith("图 5｜跨队列"):
            cap_idx[5] = (i, p)
        if p._p.findall(qn("w:r") + "/" + qn("w:drawing")) or p._p.iter(qn("w:drawing")):
            if any(True for _ in p._p.iter(qn("w:drawing"))):
                img_paras.append((i, p))
    assert set(cap_idx) == {1, 2, 3, 4, 5}, cap_idx.keys()
    print(f"图注定位: 图1-5 OK; 现有图片段落 {len(img_paras)} 个")

    # 2) 删除旧图片: 保留图 1 的图片 (= 图 1 注之前最近的那张), 其余删除
    removed = 0
    for i, p in img_paras:
        if i < cap_idx[1][0]:   # 图 1 注之前的图片 = 图 1 本体, 保留
            continue
        p._p.getparent().remove(p._p)
        removed += 1
    print(f"删除旧图片段落: {removed} 个 (保留图 1 图片)")

    # 3) 在图 2-5 图注段前插入新图 (add_picture 追加到文末后搬移到位)
    for n in (2, 3, 4, 5):
        png = str(FIGD / f"fig{n}.png")
        doc.add_picture(png, width=Cm(18))
        new_p = doc.paragraphs[-1]._p          # add_picture 生成的段落
        cap_idx[n][1]._p.addprevious(new_p)       # 搬到对应图注前
        print(f"图 {n}: 已嵌入 {png.split(chr(92))[-1] if chr(92) in png else png.split('/')[-1]} (180mm 宽)")

    doc.save(str(F))

    # 4) 复核
    d2 = docx.Document(str(F))
    shapes = d2.inline_shapes
    print(f"复核: 嵌入图片 {len(shapes)} 张")
    for sh in shapes:
        for pi, p in enumerate(d2.paragraphs):
            if sh._inline in p._p.iter():
                nxt = d2.paragraphs[pi + 1].text.strip()[:34] if pi + 1 < len(d2.paragraphs) else ""
                print(f"  {sh.width.mm:.0f}×{sh.height.mm:.0f}mm @P{pi} 下一段: {nxt!r}")
                break


if __name__ == "__main__":
    main()
