#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""V15 图注替换 + 正文图号引用插入 (2026-08-30).

输入: PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260829.docx
输出: PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260830.docx

手术内容 (全部新增文字按 0829 稿红字惯例标红):
1. 图说明区块: P116/117/118 旧图注 → fig2/3/4_caption_cn.txt 定稿图注 (S5→S5与S6 编号修正);
   P118 后插入图 5 图注
2. 正文图号引用: P80 首句末（图 2）; P84 首句末（图 3）; P88 （补充表 S5）→（图 4；补充表 S5）;
   P95 （表 1；表 3）→（表 1；表 3；图 5）
安全模式: 段落替换先清 runs 再加单 run; 正文插入按 run 拆分保格式。
"""
from __future__ import annotations

import copy
import re
from pathlib import Path

import docx
from docx.oxml.ns import qn

BASE = Path(r"F:/MIMIC3_1/V13/manuscript")
SRC = BASE / "PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260829.docx"
DST = BASE / "PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260830.docx"
FIGD = BASE / "figures"


def load_caption(n: int) -> str:
    txt = (FIGD / f"fig{n}_caption_cn.txt").read_text(encoding="utf-8")
    return "• " + "".join(line.strip() for line in txt.splitlines() if line.strip())


def red_rpr(doc) -> "ct":
    """取正文红字 run 的 rPr 作为模板 (P90 首个红 run)。"""
    for r in doc.paragraphs[90].runs:
        if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == "FF0000":
            return copy.deepcopy(r._r.find(qn("w:rPr")))
    raise RuntimeError("未找到红字模板 run")


def is_red(run) -> bool:
    return bool(run.font.color and run.font.color.rgb and str(run.font.color.rgb) == "FF0000")


def set_paragraph_text_red(p, text, rpr_tmpl):
    """清空段落 runs, 写入单红 run。"""
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    r_el = copy.deepcopy(rpr_tmpl.getparent())  # 模板 w:r
    # 重建: 新 w:r = rPr 模板 + 文本
    from docx.oxml import OxmlElement
    new_r = OxmlElement("w:r")
    new_r.append(copy.deepcopy(rpr_tmpl))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    p._p.append(new_r)


def insert_red_before_anchor(p, anchor: str, ins: str, rpr_tmpl) -> bool:
    """在段落全文中 anchor 首次出现处之前插入红色文字 ins (run 级拆分保格式)。"""
    full = p.text
    pos = full.find(anchor)
    if pos < 0:
        return False
    off = 0
    for r in p.runs:
        rl = len(r.text)
        if off <= pos < off + rl:
            local = pos - off
            if is_red(r):  # 所在 run 已红: 直接插入
                r.text = r.text[:local] + ins + r.text[local:]
                return True
            # 拆 run: 前段(原色) + 插入(红) + 后段(原色)
            from docx.oxml import OxmlElement
            r_el = r._r
            after = OxmlElement("w:r")
            if r_el.find(qn("w:rPr")) is not None:
                after.append(copy.deepcopy(r_el.find(qn("w:rPr"))))
            t_after = OxmlElement("w:t")
            t_after.set(qn("xml:space"), "preserve")
            t_after.text = r.text[local:]
            after.append(t_after)
            r_el.addnext(after)
            mid = OxmlElement("w:r")
            mid.append(copy.deepcopy(rpr_tmpl))
            t_mid = OxmlElement("w:t"); t_mid.set(qn("xml:space"), "preserve")
            t_mid.text = ins
            mid.append(t_mid)
            r_el.addnext(mid)
            r.text = r.text[:local]
            return True
        off += rl
    return False


def replace_in_paragraph(p, old: str, new: str, rpr_tmpl) -> bool:
    """段落内文本替换; 跨 run 时合并到首 run (红模板)。"""
    full = p.text
    if old not in full:
        return False
    # 定位覆盖 old 的 run 区间, 文本重排: 首 run 改写, 其余 run 删被覆盖部分
    pos = full.find(old)
    end = pos + len(old)
    off = 0
    first_done = False
    for r in list(p.runs):
        rl = len(r.text)
        rs, re_ = off, off + rl
        if re_ <= pos or rs >= end:
            off = re_
            continue
        s = max(pos - rs, 0)
        e = min(end - rs, rl)
        if not first_done:
            r.text = r.text[:s] + new + r.text[e:]
            first_done = True
        else:
            r.text = r.text[:s] + r.text[e:]
        off = re_
    return True


def main():
    doc = docx.Document(str(SRC))
    ps = doc.paragraphs
    rpr = red_rpr(doc)

    # ── 1. 图注替换 (P116/117/118) + 图 5 插入 ──
    assert ps[116].text.startswith("• 图 2｜"), ps[116].text[:20]
    assert ps[117].text.startswith("• 图 3｜"), ps[117].text[:20]
    assert ps[118].text.startswith("• 图 4｜"), ps[118].text[:20]
    cap4 = load_caption(4).replace("完整结果见补充表 S5。", "完整结果见补充表 S5 与 S6。")
    assert "S5 与 S6" in cap4
    for idx, cap in [(116, load_caption(2)), (117, load_caption(3)), (118, cap4)]:
        set_paragraph_text_red(ps[idx], cap, rpr)
    # 图 5: deepcopy 已替换的 P116 (红, 样式正确), 改文本, 插到 P118 后
    p5 = copy.deepcopy(ps[116]._p)
    ps[118]._p.addnext(p5)
    ps = doc.paragraphs  # 重新索引
    p5_para = ps[119]
    assert p5_para.text.startswith("• 图 2｜")  # deepcopy 副本
    set_paragraph_text_red(p5_para, load_caption(5), rpr)

    # ── 2. 正文图号引用 ──
    ps = doc.paragraphs
    def find_para(sub: str):
        for p in ps:
            if sub in p.text:
                return p
        raise RuntimeError(f"未找到段落: {sub[:30]}")

    p80 = find_para("逐器官分析显示，这一总体排序并不适用于所有器官")
    ok1 = insert_red_before_anchor(p80, "。", "（图 2）", rpr)  # 首个。= 首句末
    p84 = find_para("鉴于循环分量的静态状态延续误差最高")
    ok2 = insert_red_before_anchor(p84, "。", "（图 3）", rpr)
    p88 = find_para("四方面核查 DRCM（补充表 S5）")
    ok3 = replace_in_paragraph(p88, "（补充表 S5）", "（图 4；补充表 S5）", rpr)
    p95 = find_para("使用 MIMIC-IV 本地数据估计（表 1；表 3）")
    ok4 = replace_in_paragraph(p95, "（表 1；表 3）", "（表 1；表 3；图 5）", rpr)
    assert all([ok1, ok2, ok3, ok4]), [ok1, ok2, ok3, ok4]

    doc.save(str(DST))
    print(f"手术完成 → {DST.name}")
    print("图注: 图2/3/4 替换 + 图5 插入; 引用: R2/R3/R4/R5 各一处")


if __name__ == "__main__":
    main()
