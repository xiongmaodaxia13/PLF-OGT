#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""V15 冻结前一致性清理 (2026-08-30, 用户审阅意见六项).

1. 图2注删"2.6–2.8 倍"句 (两区块 + txt)
2. 图4注"表 4 注"→"补充表 S5 注" (两区块 + txt)
3. S5 注去编辑过程语言 + 口径汇总方式重写
4. Discussion 循环句: MIMIC 跨时距过度表述 → 6 h 限定表述
5. 图片移入结论后"图说明"区块, 删除 S8 后重复的第二套图注区 (P190..P204)
6. 删除孤立参考文献 25/26 (Rajendran/Lim), 27–39 重编号为 25–37, 正文引用同步 −2
所有修改按红字惯例标红。覆盖保存 20260830.docx。
"""
from __future__ import annotations

import copy
import re
import time
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(r"F:/MIMIC3_1/V13/manuscript")
F = BASE / "PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260830.docx"
FIGD = BASE / "figures"

NEW_S5_NOTE = (
    "注：患者对应关系负控的 0.567 为 3 次独立训练均值；固定模型信息移除的 0.584 为 3 次训练逐次估计的"
    "中位数；0.569 为 seed 42 单模型结构消融结果；表 2 的 0.610 为 3-seed logits ensemble 主分析结果，"
    "各口径不可直接比较。完整逐次训练、输入遮蔽、双向探针、槽匹配及 12 个代理逐项结果见补充表 S6；"
    "Shapley 逐次明细与患者特异性逐项结果见补充数据文件。"
)
DISC_OLD = ("但在循环分量中，标准 Transformer 与 PLF-OGT 均表现出跨时距稳定的重构优势，"
            "该模式在 MIMIC-IV 独立再开发中得到重现。")
DISC_NEW = ("但在 GMUICU 的循环分量中，标准 Transformer 与 PLF-OGT 均表现出跨时距稳定的重构优势；"
            "MIMIC-IV 的 6 h 分析同样观察到两种学习模型优于静态状态延续对照。")


def red_rpr(doc):
    for p in doc.paragraphs:
        if p.text.startswith("• 图 2｜总 SOFA"):
            for r in p.runs:
                if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == "FF0000":
                    return copy.deepcopy(r._r.find(qn("w:rPr")))
    raise RuntimeError("未找到红字模板")


def span_replace_red(p, old: str, new: str, rpr):
    """跨 run 替换 old → 红色新文本 (拆 run 保前后格式)。"""
    full = p.text
    pos = full.find(old)
    if pos < 0:
        return False
    end = pos + len(old)
    off = 0
    first = True
    for r in list(p.runs):
        rl = len(r.text)
        rs, re_ = off, off + rl
        if re_ <= pos or rs >= end:
            off = re_
            continue
        s, e = max(pos - rs, 0), min(end - rs, rl)
        if first:
            head, tail = r.text[:s], r.text[e:]
            r.text = head
            anchor = r._r
            if tail:
                tail_r = OxmlElement("w:r")
                if anchor.find(qn("w:rPr")) is not None:
                    tail_r.append(copy.deepcopy(anchor.find(qn("w:rPr"))))
                t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = tail
                tail_r.append(t)
                anchor.addnext(tail_r)
            mid = OxmlElement("w:r")
            mid.append(copy.deepcopy(rpr))
            t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = new
            mid.append(t)
            anchor.addnext(mid)
            first = False
        else:
            r.text = r.text[:s] + r.text[e:]
        off = re_
    return True


def set_red(p, text, rpr):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    new_r = OxmlElement("w:r")
    new_r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    new_r.append(t)
    p._p.append(new_r)


def find_all(doc, sub, start=0):
    return [p for p in doc.paragraphs[start:] if sub in p.text]


def renum_inner(inner: str) -> str:
    parts = re.split(r"([,，–\-])", inner)
    out = []
    for i, seg in enumerate(parts):
        if i % 2 == 0 and seg.isdigit():
            n = int(seg)
            out.append(str(n - 2) if n >= 27 else seg)
        else:
            out.append(seg)
    return "".join(out)


def renum_token(m):
    return "[" + renum_inner(m.group(1)) + "]"


def main():
    doc = docx.Document(str(F))
    rpr = red_rpr(doc)
    log = []

    # ── 1) 图2注 倍数句 (两区块) ──
    old21 = "预测时距保持一致；3 h 后静态状态延续误差约为学习模型的 2.6–2.8 倍。"
    n = sum(span_replace_red(p, old21, "预测时距保持一致。", rpr) for p in find_all(doc, "2.6–2.8"))
    log.append(f"1) 图2注倍数句: {n} 处")

    # ── 2) 图4注 表4注 → 补充表 S5 注 (两区块) ──
    n = sum(span_replace_red(p, "口径对照见表 4 注。", "口径对照见补充表 S5 注。", rpr)
            for p in find_all(doc, "口径对照见表 4 注"))
    log.append(f"2) 图4注表4注残留: {n} 处")

    # ── 3) S5 注重写 ──
    for p in find_all(doc, "本表由原主文表 4"):
        set_red(p, NEW_S5_NOTE, rpr)
        log.append("3) S5 注已重写")
        break

    # ── 4) Discussion 循环句 ──
    hit = False
    for p in find_all(doc, "该模式在 MIMIC-IV 独立再开发中得到重现"):
        hit = span_replace_red(p, DISC_OLD, DISC_NEW, rpr) or hit
    log.append(f"4) Discussion 跨时距句: {'已改' if hit else '!!未找到'}")

    # ── 5) 图片移入图说明区块 + 删除第二套图注区 ──
    ps = doc.paragraphs
    bullets = {}   # 图号 -> 注段落 (• 图 N｜)
    end_caps = {}  # 图号 -> 注段落 (图 N｜, 末尾区块)
    img_paras = []
    heading = refhead = None
    for i, p in enumerate(ps):
        t = p.text.strip()
        m = re.match(r"^• 图 (\d)｜", t)
        if m and i < 130:
            bullets[int(m.group(1))] = p
        m2 = re.match(r"^图 (\d)｜", t)
        if m2 and i > 185:
            end_caps[int(m2.group(1))] = p
        if t == "图（Figures）":
            heading = p
        if t.startswith("参考文献"):
            refhead = p
        if any(True for _ in p._p.iter(qn("w:drawing"))):
            img_paras.append((i, p))
    assert set(bullets) == {1, 2, 3, 4, 5} and set(end_caps) == {1, 2, 3, 4, 5}
    assert heading is not None and refhead is not None
    b1_pos = ps.index(bullets[1])
    for i, p in img_paras:  # 图1图片=末尾区块图1注之前那张; 按出现顺序对应图1-5
        pass
    order_imgs = [p for _, p in img_paras]
    # 末尾区块顺序: img(图1) cap1 img(图2) cap2 ... 依据 dom 顺序配对
    seq = sorted(img_paras + [(ps.index(c), c) for c in end_caps.values()], key=lambda x: x[0])
    img_of = {}
    last_img = None
    for idx, p in seq:
        if any(True for _ in p._p.iter(qn("w:drawing"))):
            last_img = p
        else:
            m = re.match(r"^图 (\d)｜", p.text.strip())
            if m and last_img is not None:
                img_of[int(m.group(1))] = last_img
                last_img = None
    assert set(img_of) == {1, 2, 3, 4, 5}, img_of.keys()
    for nfig in (1, 2, 3, 4, 5):
        bullets[nfig]._p.addprevious(img_of[nfig]._p)  # 移动图片元素到 • 图N 注之前
    # 删除第二套区块: heading 至 refhead 之间的所有段落
    node = heading._p.getnext()
    removed = 0
    while node is not None and node is not refhead._p:
        nxt = node.getnext()
        node.getparent().remove(node)
        removed += 1
        node = nxt
    heading._p.getparent().remove(heading._p)
    log.append(f"5) 图片 5 张移入图说明区块; 删除末尾图区段落 {removed + 1} 个")

    # ── 6) 孤立文献 25/26 删除 + 27–39 重编号 + 正文引用 −2 ──
    ps = doc.paragraphs
    for p in list(ps):
        t = p.text.strip()
        if (t.startswith("25.") and "Rajendran" in t) or (t.startswith("26.") and "Lim" in t):
            p._p.getparent().remove(p._p)
    nref = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r"^(3[0-9]|2[7-9])\.", t)
        if m:
            newn = int(m.group(1)) - 2
            span_replace_red(p, m.group(1) + ".", f"{newn}.", rpr)
            nref += 1
    ncit = 0
    pat = re.compile(r"\[([0-9]{1,2}(?:[,，–\-][0-9]{1,2})*)\]")
    for p in doc.paragraphs:
        toks = {m.group(0) for m in pat.finditer(p.text)
                if any(int(x) >= 27 for x in re.findall(r"\d+", m.group(1)))}
        for tok in toks:
            inner = tok[1:-1]
            newtok = "[" + renum_inner(inner) + "]"
            if newtok == tok:
                continue
            while span_replace_red(p, tok, newtok, rpr):
                ncit += 1
    log.append(f"6) 删文献 25/26; 重编号文献 {nref} 条; 正文引用改 {ncit} 处")

    # ── 保存 (锁定重试) ──
    for attempt in range(3):
        try:
            doc.save(str(F))
            print("\n".join(log))
            print("已保存:", F.name)
            return
        except PermissionError:
            time.sleep(8)
    doc.save(str(F).replace(".docx", "_new.docx"))
    print("\n".join(log))
    print("!! 原文件被占用, 已保存为 _new.docx — 关闭后重命名即可")


if __name__ == "__main__":
    main()
