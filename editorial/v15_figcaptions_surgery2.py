#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""V15 图注手术第 2 处: 文档末尾完整图注区块 (P192-198 区域) 的旧图 2/3/4 替换 + 图 5 插入.

0829/0830 稿有两个图注区: P114-119 项目符号版 (第 1 次手术已换) 与文档末尾完整版 (本脚本)。
"""
from __future__ import annotations

import copy
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(r"F:/MIMIC3_1/V13/manuscript")
F = BASE / "PLF_OGT_V15_图表分工与补充表重组_红字审核版_20260830.docx"
FIGD = BASE / "figures"

OLD_PREFIX = {
    2: "图 2｜TCR 的事实治疗路径轨迹重放",
    3: "图 3｜事实治疗路径信息的功能利用",
    4: "图 4｜残差状态的患者特异性检验",
}


def load_caption(n: int, bullet: bool = False) -> str:
    txt = (FIGD / f"fig{n}_caption_cn.txt").read_text(encoding="utf-8")
    body = "".join(line.strip() for line in txt.splitlines() if line.strip())
    return ("• " + body) if bullet else body


def red_rpr(doc):
    for p in doc.paragraphs:
        if p.text.startswith("• 图 2｜总 SOFA"):
            for r in p.runs:
                if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == "FF0000":
                    return copy.deepcopy(r._r.find(qn("w:rPr")))
    raise RuntimeError("未找到红字模板")


def set_paragraph_text_red(p, text, rpr_tmpl):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    new_r = OxmlElement("w:r")
    new_r.append(copy.deepcopy(rpr_tmpl))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    p._p.append(new_r)


def main():
    doc = docx.Document(str(F))
    rpr = red_rpr(doc)
    ps = doc.paragraphs

    cap4 = load_caption(4).replace("完整结果见补充表 S5。", "完整结果见补充表 S5 与 S6。")
    targets = {}
    for i, p in enumerate(ps):
        t = p.text.strip()
        if i > 150:  # 只处理文档末尾区块
            for n, pref in OLD_PREFIX.items():
                if t.startswith(pref):
                    targets[n] = i
    assert set(targets) == {2, 3, 4}, targets

    for n in (2, 3, 4):
        set_paragraph_text_red(ps[targets[n]], load_caption(n) if n != 4 else cap4, rpr)
    # 图 5 插到 图 4 段后
    p5 = copy.deepcopy(ps[targets[4]]._p)
    ps[targets[4]]._p.addnext(p5)
    ps2 = doc.paragraphs
    p5_para = next(p for p in ps2 if p._p is p5)
    set_paragraph_text_red(p5_para, load_caption(5), rpr)

    doc.save(str(F))
    print("第 2 处图注区手术完成: 末尾区块 图2/3/4 替换 + 图5 插入")


if __name__ == "__main__":
    main()
