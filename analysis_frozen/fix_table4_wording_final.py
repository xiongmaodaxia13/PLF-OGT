#!/usr/bin/env python
"""修 Table 4 列错位 + R5/综合段措辞 + Figure 4 图注."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from pathlib import Path

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_p0_final.docx"
out = D / "方法+结果0808_frozen_p0_v2.docx"

d = Document(str(src))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# P0-1: Table 4 trajectory 行 — 修正列错位 + 补 GMUICU 对照值
# ============================================================
# 表3 (表头: 指标 | GMUICU | MIMIC-IV) 是 Table 4
# 新增的 6 行 trajectory 数据被放到了 GMUICU 列(列1), 应该是:
# GMUICU列放GMUICU值, MIMIC列放MIMIC值
# GMUICU frozen 值: PLF 1.002, persist 0.645, TR 0.725, 69.2%, changed PLF 1.332, changed TR 1.145, changed persist 2.018
# MIMIC frozen 值: PLF 0.691, persist 0.604, TR 1.223, 60.2%, changed PLF 1.085, changed TR 1.419, changed persist 1.443

t4 = None
for ti, t in enumerate(d.tables):
    hdr = t.rows[0].cells[0].text if t.rows else ""
    if "指标" in hdr and "GMUICU" in t.rows[0].cells[1].text:
        t4 = t
        break

if t4:
    # 找 trajectory 行 (后6行)
    trajectory_data = [
        # (行关键词, GMUICU值, MIMIC值)
        ("PLF trajectory MAE", "1.002", "0.691"),
        ("persistence MAE", "0.645", "0.604"),
        ("Transformer MAE", "0.725", "1.223"),
        ("ΔSOFA=0", "69.2", "60.2"),
        ("PLF MAE", "1.332", "1.085"),  # |ΔSOFA|≥1
        ("TR MAE", "1.145", "1.419"),   # |ΔSOFA|≥1
    ]
    for ri in range(len(t4.rows) - 6, len(t4.rows)):
        row = t4.rows[ri]
        cell0_text = row.cells[0].text
        for keyword, gm_val, mim_val in trajectory_data:
            if keyword in cell0_text:
                # 列1 = GMUICU, 列2 = MIMIC-IV
                row.cells[1].paragraphs[0].runs[0].text = gm_val if row.cells[1].paragraphs[0].runs else gm_val
                row.cells[2].paragraphs[0].runs[0].text = mim_val if row.cells[2].paragraphs[0].runs else mim_val
                rev.append(f"Table4 行{ri}: {keyword} GMUICU={gm_val} MIMIC={mim_val}")
                break

# ============================================================
# P0-2: R5段97 — 措辞改为"队列依赖性"而非固定trade-off
# ============================================================
p97 = d.paragraphs[97]
# 替换 "performance–auditability trade-off 的具体幅度可能受队列事件率和训练规模影响"
rep_runs(p97,
    "提示 performance–auditability trade-off 的具体幅度可能受队列事件率和训练规模影响。",
    "提示两种模型的相对预测性能具有队列依赖性。")
rev.append("R5段97: trade-off措辞→队列依赖性")

# ============================================================
# P0-3: 综合段(段106) — 调整为"队列依赖性"定位
# ============================================================
p106 = d.paragraphs[106]
rep_runs(p106,
    "尽管 Transformer 的轨迹误差仍低于 PLF-OGT。",
    "尽管 Transformer 的轨迹误差在 GMUICU 中低于 PLF-OGT，而在 MIMIC-IV 中排序反转（见 Result 5）。")
rep_runs(p106,
    "PLF-OGT 的主要增量在于 structured auditability 和 patient-specific predictive representation，"
    "而非相对于标准 Transformer 的预测精度优势。",
    "PLF-OGT 未显示相对于标准 Transformer 的跨队列一致预测精度优势："
    "Transformer 在 GMUICU 中具有更低轨迹误差，而 PLF-OGT 在 MIMIC-IV changed-state 中具有更低误差。"
    "相比之下，PLF-OGT 更稳定的设计增量在于 structured auditability 和 patient-specific predictive representation。")
rev.append("综合段: 调整为队列依赖性定位")

# ============================================================
# P1-1: Figure 4 图注修正
# ============================================================
for i, p in enumerate(d.paragraphs):
    if "patient-specificity 审计" in p.text and "代理恢复" in p.text:
        rep_runs(p,
            "代理恢复与随机映射 + S/R 2×2 + patient-specificity negative controls + 成功槽与失败槽。",
            "(a) 6 h AUPRC under matched, population-mean, query-only and stay-level shuffled residual states. "
            "(b) Corresponding macro-MAE. Error bars indicate 3-seed SD.")
        rev.append(f"Figure 4 段{i}: 图注修正")
    elif "Patient-specificity 审计" in p.text and "代理恢复" in p.text:
        rep_runs(p,
            "代理恢复与随机映射 + S/R 2×2 + patient-specificity negative controls + 成功槽与失败槽。",
            "(a) 6 h AUPRC under matched, population-mean, query-only and stay-level shuffled residual states. "
            "(b) Corresponding macro-MAE. Error bars indicate 3-seed SD.")
        rev.append(f"Figure 4 段{i}: 图注修正(变体)")

# ============================================================
# P1-2: R5 "部分方向性判别" 旧句升级
# ============================================================
for i, p in enumerate(d.paragraphs):
    if "部分方向性判别结果与代理对齐" in p.text:
        rep_runs(p,
            "MIMIC-IV结果仅支持……部分方向性判别结果与代理对齐特征，而不支持完整多器官轨迹性能",
            "MIMIC-IV 独立再开发重现了 TCR–care-off 的方向性判别差异、完整 S+R 通路的优势，"
            "以及 6 h changed-state trajectory sensitivity 的方向")
        rev.append(f"段{i}: R5旧句升级")
        break

d.save(str(out))
print("修订:")
for r in rev: print(f"  {r}")
print(f"保存: {out}")
