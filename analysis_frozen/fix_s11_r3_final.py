#!/usr/bin/env python
"""修正 S11 表格错位 + R3 段82 合并 + R5 补句."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
import json
from pathlib import Path

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_r3.docx"
out = D / "方法+结果0808_frozen_final.docx"
RD = D.parent / "results/v4"

d = Document(str(src))
org = json.load(open(str(RD / "frozen_organ_noncv.json")))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# 1. S11 (表14) 彻底重建 — 从 frozen JSON 正确填入
# ============================================================
# frozen_organ_noncv.json organ_mae:
#   Resp:  tcr=0.815 co=0.822 delta=+0.006
#   Cv:    tcr=0.177 co=0.428 delta=+0.251
#   Renal: tcr=0.341 co=0.345 delta=+0.003
#   Coag:  tcr=0.131 co=0.130 delta=-0.001
#   Hepatic: tcr=0.217 co=0.218 delta=+0.001
#   CNS:   tcr=0.147 co=0.199 delta=+0.053
#
# 表14 当前7行x4列, 行0=表头, 行1-6=器官
# 但器官名和数值错位: 行1器官名=呼吸, 但行1数值=0.815/0.818(呼吸TCR/旧co)
# 实际错位: 器官名和数值差了一行
#
# 彻底修复: 逐行设置全部单元格

ORGAN_CN = ["呼吸", "心血管", "肾脏", "凝血", "肝脏", "中枢神经"]
ORGAN_KEY = ["Resp", "Cv", "Renal", "Coag", "Hepatic", "CNS"]

t14 = d.tables[14]  # S11

# 表头 (行0): 统一符号定义
# Δ = care-off MAE − TCR MAE (正值=care-off更差)
t14.rows[0].cells[0].paragraphs[0].runs[0].text = "器官"
t14.rows[0].cells[1].paragraphs[0].runs[0].text = "TCR 6h MAE"
t14.rows[0].cells[2].paragraphs[0].runs[0].text = "care-off 6h MAE"
t14.rows[0].cells[3].paragraphs[0].runs[0].text = "Δ（care-off − TCR）"

for i, (cn, key) in enumerate(zip(ORGAN_CN, ORGAN_KEY)):
    r = org["organ_mae"][key]
    row = t14.rows[i + 1]
    row.cells[0].paragraphs[0].runs[0].text = cn
    row.cells[1].paragraphs[0].runs[0].text = f"{r['tcr_mae']:.3f}"
    row.cells[2].paragraphs[0].runs[0].text = f"{r['co_mae']:.3f}"
    row.cells[3].paragraphs[0].runs[0].text = f"{r['delta']:+.3f}"

rev.append("1. S11(表14) 彻底重建 — 6器官从frozen JSON正确填入")

# ============================================================
# 2. R3 段82: 合并到段81, 段82清空
#    段81已经是frozen新版, 段82是旧版重复(含0.200)
#    做法: 段82重写为审稿人建议的精简单段(含ΔBrier)
#    段81保留当前内容但删除最后一句"信息轴"
# ============================================================

# 段81: 删除"信息轴"句
p81 = d.paragraphs[81]
rep_runs(p81,
    "当前证据支持一条以心血管为主轴的治疗—器官信息轴，而非治疗信息在多器官系统间的广泛传播。",
    "这些结果提示模型对锚点后治疗信息的功能依赖主要由心血管分量驱动，而不支持治疗信息在六个器官系统间产生广泛且均一的预测增益。")
rev.append("2a. R3段81 删除'信息轴'→更保守表述")

# 段82: 清空(合并到段81)
set_para(d.paragraphs[82], "")
rev.append("2b. R3段82 清空(合并到段81)")

# ============================================================
# 3. R5 正文补句: S-only/R-only 未复制
#    在段96 找到 S/R 2×2 数据后补一句
# ============================================================
p96 = d.paragraphs[96]
# 检查是否已有"未复制"
if "未复制" not in p96.text and "近似" not in p96.text:
    # 在 S/R 2×2 数据句后追加
    rep_runs(p96,
        "φ_R/φ_S Brier 效用 3/3 seed 成立。",
        "φ_R/φ_S Brier 效用 3/3 seed 成立。与 GMUICU 不同，MIMIC-IV 中 S-only 与 R-only 的性能近似（0.178 vs 0.177），未复制 GMUICU 中 residual-only 高于 concept-only 的相对排序。")
    rev.append("3. R5段96 补S/R排序未复制句")

# ============================================================
# 4. Methods "排序一致性" → "通路阻断结果"
# ============================================================
for i, p in enumerate(d.paragraphs):
    if "排序一致性" in p.text:
        rep_runs(p, "排序一致性", "通路阻断结果")
        rev.append(f"4. Methods段{i} 排序一致性→通路阻断结果")
        break

d.save(str(out))
print("修订完成:")
for r in rev: print(f"  {r}")
print(f"\n保存: {out}")
