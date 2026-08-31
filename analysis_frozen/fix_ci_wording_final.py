#!/usr/bin/env python
"""写入 changed-state paired CI + 措辞升级."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from pathlib import Path

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_p0_v2.docx"
out = D / "方法+结果0808_frozen_ci.docx"

d = Document(str(src))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# 1. R2段76: changed-state 补 paired CI
# ============================================================
p76 = d.paragraphs[76]
# 在 "由此，两种深度学习模型在实际发生临床可观察状态转变的锚点中均明显优于静态 persistence 假设"
# 后补 CI
rep_runs(p76,
    "由此，两种深度学习模型在实际发生临床可观察状态转变的锚点中均明显优于静态 persistence 假设，而标准 Transformer 的轨迹误差仍低于 PLF-OGT。",
    "ICU-stay 级配对 cluster bootstrap（2,000 次）显示，在 |ΔSOFA|≥1 锚点中 PLF-OGT 和 Transformer 相对 persistence 的 SOFA MAE 优势分别为 −0.685（95% CI −0.752 至 −0.617）和 −0.873（−0.939 至 −0.809），均不跨零；PLF-OGT 相对 Transformer 的 MAE 差值为 +0.188（+0.148 至 +0.231），亦不跨零，表明在 GMUICU 中 Transformer 的轨迹精度显著高于 PLF-OGT。")
rev.append("1. R2段76: 补changed-state paired CI (GMUICU)")

# ============================================================
# 2. R5段97: MIMIC TR comparator 补 CI + 措辞升级
# ============================================================
p97 = d.paragraphs[97]
# 替换 "PLF-OGT 在此子集中显示低于 Transformer 的轨迹误差" → 加 CI
rep_runs(p97,
    "PLF-OGT 在此子集中显示低于 Transformer 的轨迹误差。",
    "ICU-stay 级配对 bootstrap 显示，PLF-OGT 相对 Transformer 的 MAE 优势在 MIMIC-IV 全部锚点为 −0.531（95% CI −0.570 至 −0.496），在 |ΔSOFA|≥1 锚点为 −0.333（−0.383 至 −0.286），均不跨零且方向与 GMUICU 相反，确认两种模型的相对轨迹精度排序具有队列依赖性。")
# 删除 "该模式与 GMUICU 不完全一致...队列依赖性" (已被上面CI替代)
rep_runs(p97,
    "该模式与 GMUICU 不完全一致（GMUICU 中 Transformer 在变化锚点仍优于 PLF-OGT），提示两种模型的相对预测性能具有队列依赖性。",
    "")
rev.append("2. R5段97: MIMIC TR CI + 措辞升级")

# ============================================================
# 3. 综合段106: 措辞精确化
# ============================================================
p106 = d.paragraphs[106]
rep_runs(p106,
    "Transformer 在 GMUICU 中具有更低轨迹误差，而 PLF-OGT 在 MIMIC-IV changed-state 中具有更低误差。",
    "Transformer 在 GMUICU 的总体及状态变化锚点中具有更低轨迹误差（paired bootstrap CI 不跨零），而 PLF-OGT 在 MIMIC-IV 的总体及状态变化锚点中具有更低点估计（paired bootstrap CI 不跨零）。")
rev.append("3. 综合段: CI措辞精确化(两队列都不跨零)")

# ============================================================
# 4. Table 4 补 CI 行 (可选, 审稿人推荐)
# ============================================================
t4 = None
for ti, t in enumerate(d.tables):
    hdr = t.rows[0].cells[0].text if t.rows else ""
    if "指标" in hdr and "GMUICU" in t.rows[0].cells[1].text:
        t4 = t; break

if t4:
    # 在 changed-state 行后补 paired CI 行
    ci_rows = [
        ("PLF−TR CI (all)", "+0.276 (+0.235,+0.317)", "−0.531 (−0.570,−0.496)"),
        ("PLF−TR CI (|ΔSOFA|≥1)", "+0.188 (+0.148,+0.231)", "−0.333 (−0.383,−0.286)"),
        ("PLF−persist CI (|ΔSOFA|≥1)", "−0.685 (−0.752,−0.617)", "−0.358 (−0.378,−0.338)"),
    ]
    for label, gm, mim in ci_rows:
        row = t4.add_row()
        row.cells[0].paragraphs[0].add_run(label)
        row.cells[1].paragraphs[0].add_run(gm)
        row.cells[2].paragraphs[0].add_run(mim)
    rev.append("4. Table 4 补3行paired CI")

d.save(str(out))
print("修订:")
for r in rev: print(f"  {r}")
print(f"保存: {out}")
