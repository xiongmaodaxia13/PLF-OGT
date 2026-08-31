#!/usr/bin/env python
"""Apply the final frozen consistency fixes to 方法+结果0808_frozen_final.docx.

Fixes:
  1. S3 refill: TRUE macro-MAE + sofa MAE with CIs (drop cumulative column)
  2. Table 2 / S2: OLP rows to frozen OLP values; 6h OLP CI replaced; delta rows 1/3/12h
  3. S11: sign note + caption moved above table + keep-with-next
  4. Result 5: add 'not replicated' sentence + S12 pointer
  5. New S12 (MIMIC multi-horizon discrimination + care-off paired)
  6. Supplementary list: dedupe + reorder + add S12
  7. Result 3 macro numbers; Result 1 event counts; Methods macro definition
"""
import copy
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.table import Table, _Row
from docx.oxml.ns import qn

SRC = r"F:/MIMIC3_1/V13/manuscript/方法+结果0808_frozen_final.docx"
OUT = r"F:/MIMIC3_1/V13/manuscript/方法+结果0808_frozen_final_修订.docx"
CI_JSON = Path(r"F:/MIMIC3_1/V13/results/v4/recompute_s3_full_ci.json")
MIMIC_JSON = Path(r"F:/MIMIC3_1/V13/results_mimic/mimic_3seed_corrected.json")

ci = json.loads(CI_JSON.read_text(encoding="utf-8"))
mimic = json.loads(MIMIC_JSON.read_text(encoding="utf-8"))


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError("paragraph not found: " + prefix)


def find_table(doc, header_first):
    for t in doc.tables:
        if t.rows and t.rows[0].cells[0].text.strip() == header_first:
            return t
    raise RuntimeError("table not found: " + header_first)


def set_cell_text(cell, text):
    for pp in cell.paragraphs[1:]:
        pp._element.getparent().remove(pp._element)
    p = cell.paragraphs[0]
    rpr = p.runs[0]._element.rPr if p.runs and p.runs[0]._element.rPr is not None else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    nr = p.add_run(text)
    if rpr is not None:
        nr._element.insert(0, copy.deepcopy(rpr))


def replace_run0(p, old, new):
    assert old in p.runs[0].text, f"anchor missing in [{p.text[:40]}]: {old[:50]}"
    p.runs[0].text = p.runs[0].text.replace(old, new)


def append_run(p, text):
    src = p.runs[-1] if p.runs else None
    r = p.add_run(text)
    if src is not None and src._element.rPr is not None:
        r._element.insert(0, copy.deepcopy(src._element.rPr))


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def fmt3(x):
    return f"{x:.3f}"


def fmt_ci(lo, hi):
    return f"{lo:.3f}–{hi:.3f}"


def fmt_delta(x):
    return f"{x:+.3f}"


def main():
    doc = Document(SRC)

    # ---------- 1) S3 refill (drop cumulative column, fill CIs) ----------
    s3 = find_table(doc, "时距")
    assert len(s3.rows[0].cells) == 8
    # remove 4th column (index 3) from every row
    for row in s3.rows:
        tc = row._tr.tc_lst[3]
        tc.getparent().remove(tc)
    grid = s3._tbl.tblGrid
    cols = grid.findall(qn("w:gridCol"))
    if len(cols) == 8:
        grid.remove(cols[3])
    # header
    header = [
        "时距",
        "TCR 总分 MAE（95% CI）",
        "TCR macro-MAE（95% CI）",
        "care-off 总分 MAE（95% CI）",
        "care-off macro-MAE（95% CI）",
        "Δ 总分（TCR−care-off，95% CI）",
        "Δ macro-MAE（TCR−care-off，95% CI）",
    ]
    for j, h in enumerate(header):
        set_cell_text(s3.rows[0].cells[j], h)
    sofa_by_h = {r["horizon"]: r for r in ci["sofa"]}
    macro_by_h = {r["horizon"]: r for r in ci["macro"]}
    for i, h in enumerate([1, 3, 6, 12], start=1):
        s = sofa_by_h[h]
        m = macro_by_h[h]
        row = [
            f"{h} h",
            f"{fmt3(s['sofa_tcr'])}（{fmt_ci(*s['sofa_tcr_ci'])}）",
            f"{fmt3(m['macro_tcr'])}（{fmt_ci(*m['macro_tcr_ci'])}）",
            f"{fmt3(s['sofa_co'])}（{fmt_ci(*s['sofa_co_ci'])}）",
            f"{fmt3(m['macro_co'])}（{fmt_ci(*m['macro_co_ci'])}）",
            f"{fmt_delta(s['delta_sofa'])}（{fmt_ci(*s['delta_sofa_ci'])}）",
            f"{fmt_delta(m['delta_macro'])}（{fmt_ci(*m['delta_macro_ci'])}）",
        ]
        for j, v in enumerate(row):
            set_cell_text(s3.rows[i].cells[j], v)
    p_s3note = find_para(doc, "注：CI 为 ICU-stay 级 cluster bootstrap")
    replace_run0(
        p_s3note,
        p_s3note.runs[0].text,
        "注：TCR 与 care-off 均为 3 次独立训练集成的点估计；CI 为 ICU-stay 级 cluster bootstrap（n_boot=2000）。"
        "macro-MAE 定义为六个器官各自全局分子/分母聚合 MAE 的算术平均；Δ = TCR − care-off。",
    )
    print("S3 refilled", flush=True)

    # ---------- 2) Table 2 main: OLP row + trajectory delta rounding ----------
    t2 = find_table(doc, "设定 / 时距")
    olp_new = {1: ("0.815", "0.119"), 3: ("0.830", "0.245"), 6: ("0.794", "0.238"), 12: ("0.724", "0.171")}
    # row 2 = PLF-OGT (OLP); columns 1..4 = 1/3/6/12 h as "AUROC / AUPRC"
    for j, h in enumerate([1, 3, 6, 12], start=1):
        set_cell_text(t2.rows[2].cells[j], f"{olp_new[h][0]} / {olp_new[h][1]}")
    # trajectory delta row (index 8): 3h -0.203 -> -0.202; 12h -0.214 -> -0.213
    set_cell_text(t2.rows[8].cells[2], "-0.202")
    set_cell_text(t2.rows[8].cells[4], "-0.213")
    print("Table 2 updated", flush=True)

    # ---------- 3) S2: OLP rows + delta rows ----------
    s2 = find_table(doc, "模型 / 设定")
    disc_by_h = {r["horizon"]: r for r in ci["discrimination"]}
    # row mapping: 1 TCR 1h, 2 TCR 3h, 3 TCR 6h, 4 TCR 12h, 5 OLP 1h ... 8 OLP 12h,
    # 9..12 delta
    for i, h in enumerate([1, 3, 6, 12], start=5):
        d = disc_by_h[h]
        set_cell_text(s2.rows[i].cells[0], "PLF-OGT（OLP）")
        set_cell_text(s2.rows[i].cells[1], f"{h}h")
        set_cell_text(s2.rows[i].cells[2], f"{fmt3(d['co_auroc'])}（{fmt_ci(*d['co_auroc_ci'])}）")
        set_cell_text(s2.rows[i].cells[3], f"{fmt3(d['co_auprc'])}（{fmt_ci(*d['co_auprc_ci'])}）")
    for i, h in enumerate([1, 3, 12], start=9):
        d = disc_by_h[h]
        set_cell_text(s2.rows[i].cells[0], "Δ（TCR−OLP，配对）")
        set_cell_text(s2.rows[i].cells[1], f"{h}h")
        set_cell_text(s2.rows[i].cells[2], f"{fmt_delta(d['delta_auroc'])}（{fmt_ci(*d['delta_auroc_ci'])}）")
        set_cell_text(s2.rows[i].cells[3], f"{fmt_delta(d['delta_auprc'])}（{fmt_ci(*d['delta_auprc_ci'])}）")
    print("S2 updated", flush=True)

    # ---------- 4) S11 caption move + note; delete misplaced S10 note ----------
    s11 = find_table(doc, "器官")
    s11_caption = find_para(doc, "补充表 S11｜六器官 6h 轨迹 MAE 与差值")
    s11_caption._element.getparent().remove(s11_caption._element)
    s11._tbl.addprevious(s11_caption._element)
    s11_caption.paragraph_format.keep_with_next = True
    # delete misplaced note under S10
    for p in doc.paragraphs:
        if p.text.strip().startswith("注：TCR 与 OLP（care-off）为 3 次独立训练 logit 集成点估计"):
            delete_paragraph(p)
            break
    print("S11 caption moved, misplaced note deleted", flush=True)

    # ---------- 5) Result 3 macro numbers ----------
    p_r3 = find_para(doc, "在固定模型参数下关闭治疗更新分支后")
    replace_run0(
        p_r3,
        "器官 macro-MAE 从 0.264 增至 0.318，对应配对差为 -0.054。",
        "器官 macro-MAE 从 0.305 增至 0.357，对应配对差为 −0.052（95% CI −0.060 至 −0.045）。",
    )
    print("Result 3 macro updated", flush=True)

    # ---------- 6) Result 1 event counts ----------
    p_r1 = find_para(doc, "GMUICU 队列共筛选 3,850 例患者")
    replace_run0(
        p_r1,
        "[待补：各预测时距的可评价锚点数及事件率，如作为正文报告]",
        "GMUICU 各时距可评价锚点为 66,619/65,867/64,751/62,563 个，1、3、6、12 h SOFA 恶化率分别为 "
        "3.05%、6.08%、7.03% 和 7.74%；MIMIC-IV 各时距可评价锚点为 43,374/43,084/42,623/41,753 个，"
        "恶化率分别为 1.74%、3.06%、4.34% 和 5.92%（MIMIC-IV 结果详见 Result 5）。",
    )
    print("Result 1 event counts filled", flush=True)

    # ---------- 7) Methods macro definition ----------
    p_m = find_para(doc, "TCR 为主要任务，完整评价未来 1、3、6 和 12 小时的连续器官轨迹")
    replace_run0(
        p_m,
        "预设主要轨迹指标为 6 小时器官 macro-MAE，主要临床量纲指标为 6 小时 SOFA 总分 MAE；六器官分量 MAE 和累计轨迹误差为次要指标。",
        "预设主要轨迹指标为 6 小时器官 macro-MAE（定义为六个器官各自全局分子/分母聚合 MAE 的算术平均），"
        "主要临床量纲指标为 6 小时 SOFA 总分 MAE；六器官分量 MAE 为次要指标。",
    )
    print("Methods macro definition updated", flush=True)

    # ---------- 8) Result 5: not-replicated sentence + S12 pointer ----------
    p_r5 = find_para(doc, "在 MIMIC-IV 三次独立再开发中")
    t = p_r5.runs[0].text
    if "未复制" not in t:
        t = t.replace(
            "完整 S+R 通路优于通路阻断条件。",
            "完整 S+R 通路优于通路阻断条件。与 GMUICU 不同，MIMIC-IV 中 S-only 与 R-only 的性能近似"
            "（0.178 vs 0.177），未复制 GMUICU 中 residual-only 高于 concept-only 的相对排序。",
        )
        if not t.endswith("。") or not t.rstrip().endswith("补充表 S12。"):
            t = t.rstrip() + "各时距完整 AUROC/AUPRC 及 95% CI、care-off 配对差值见补充表 S12。"
        p_r5.runs[0].text = t
    print("Result 5 updated", flush=True)

    # ---------- 9) Supplementary list dedupe + reorder + S12 ----------
    list_paras = []
    started = False
    for p in doc.paragraphs:
        ts = p.text.strip()
        if "补充表 S1：完整 43 变量清单" in ts:
            started = True
        if started and "• 补充表" in ts:
            list_paras.append(p)
        if started and ts == "表格":
            break
    new_list = [
        "• 补充表 S1：完整 43 变量清单（缺失率、标准化参数中位数与 IQR）",
        "• 补充表 S2–S3：多时距判别与轨迹重放完整结果（含 95% CI 与 Transformer/GRU 参考基线）",
        "• 补充表 S4：结构消融完整结果（含 3 次独立训练验证）",
        "• 补充表 S5：代理注意力富集（跨 3 次独立训练，含灌注轴 0.000×3）",
        "• 补充表 S6：表示重叠、双向探针、输入遮蔽与槽匹配",
        "• 补充表 S7：校准与阈值完整结果",
        "• 补充表 S8：推理时延与参数量",
        "• 补充表 S9：7 例/20 锚点病例级解释审计",
        "• 补充表 S10：逐次独立训练明细（S/R 2×2、Shapley、富集）",
        "• 补充表 S11：六器官 6h 轨迹 MAE 与差值（GMICU）",
        "• 补充表 S12：MIMIC-IV 多时距判别与 care-off 配对（独立再开发）",
    ]
    assert len(list_paras) >= 11, f"supp list paragraphs found: {len(list_paras)}"
    for i, text in enumerate(new_list):
        replace_run0(list_paras[i], list_paras[i].runs[0].text, text)
    for extra in list_paras[len(new_list):]:
        delete_paragraph(extra)
    print("supplementary list deduped", flush=True)

    # ---------- 10) S11 note + new S12 table/caption/note ----------
    anchor = find_para(doc, "图（Figures）")
    # S11 note (below S11 table)
    p_note = doc.add_paragraph(
        "注：ΔMAE = care-off − TCR，正值表示关闭治疗修正分支后轨迹误差增加；器官 MAE 为逐器官全局分子/分母聚合。"
    )
    anchor._element.addprevious(p_note._element)
    # S12 caption
    p_cap = doc.add_paragraph("补充表 S12｜MIMIC-IV 多时距判别与 care-off 配对（独立再开发）")
    p_cap.paragraph_format.keep_with_next = True
    anchor._element.addprevious(p_cap._element)
    # S12 table (clone S2 structure)
    s2_tbl = s2._tbl
    new_tbl_el = copy.deepcopy(s2_tbl)
    trs = new_tbl_el.findall(qn("w:tr"))
    for tr in trs[13:]:
        new_tbl_el.remove(tr)
    anchor._element.addprevious(new_tbl_el)
    s12 = Table(new_tbl_el, doc)
    s12_header = ["模型 / 设定", "时距", "AUROC（95% CI）", "AUPRC（95% CI）"]
    for j, htxt in enumerate(s12_header):
        set_cell_text(s12.rows[0].cells[j], htxt)
    m_disc = mimic["discrimination"]
    m_co = mimic["careoff_paired"]
    r = 1
    for h in [1, 3, 6, 12]:
        d = m_disc["TCR"][f"{h}h"]
        vals = [
            "PLF-OGT（TCR）", f"{h}h",
            f"{fmt3(d['auroc'])}（{fmt_ci(*d['auroc_ci'])}）",
            f"{fmt3(d['auprc'])}（{fmt_ci(*d['auprc_ci'])}）",
        ]
        for j, v in enumerate(vals):
            set_cell_text(s12.rows[r].cells[j], v)
        r += 1
    for h in [1, 3, 6, 12]:
        d = m_disc["OLP"][f"{h}h"]
        vals = [
            "PLF-OGT（OLP/care-off）", f"{h}h",
            f"{fmt3(d['auroc'])}（{fmt_ci(*d['auroc_ci'])}）",
            f"{fmt3(d['auprc'])}（{fmt_ci(*d['auprc_ci'])}）",
        ]
        for j, v in enumerate(vals):
            set_cell_text(s12.rows[r].cells[j], v)
        r += 1
    for h in [1, 3, 6, 12]:
        d = m_co[f"{h}h"]
        vals = [
            "Δ（TCR−OLP，配对）", f"{h}h",
            f"{fmt_delta(d['delta_auroc'])}（{fmt_ci(*d['delta_auroc_ci'])}）",
            f"{fmt_delta(d['delta_auprc'])}（{fmt_ci(*d['delta_auprc_ci'])}）",
        ]
        for j, v in enumerate(vals):
            set_cell_text(s12.rows[r].cells[j], v)
        r += 1
    # S12 note
    p_note2 = doc.add_paragraph(
        "注：TCR 与 OLP 均为 MIMIC-IV 三次独立再开发 logit 集成的点估计；CI 为 ICU-stay 级配对 cluster bootstrap"
        "（n_boot=2000）。各时距可评价锚点数为 43,374/43,084/42,623/41,753，6 h SOFA 恶化率 4.34%"
        "（1/3/12 h 分别为 1.74%/3.06%/5.92%）。Δ = TCR − OLP（care-off）。MIMIC-IV 为本地独立再开发模型，"
        "非冻结权重外部迁移。"
    )
    anchor._element.addprevious(p_note2._element)
    print("S11 note + S12 added", flush=True)

    # ---------- save ----------
    try:
        doc.save(SRC)
        print("SAVED ->", SRC, flush=True)
    except PermissionError:
        doc.save(OUT)
        print("SAVED ->", OUT, flush=True)


if __name__ == "__main__":
    main()
