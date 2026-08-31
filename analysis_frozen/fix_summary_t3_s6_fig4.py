#!/usr/bin/env python
"""综合段更新 + Table 3 补负控 + S6 补负控 + Figure 4 替换."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from pathlib import Path
import zipfile, json

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_r4.docx"
out = D / "方法+结果0808_frozen_r4_final.docx"
RD = D.parent / "results/v4"
FIG = D.parent / "figures/图4_概念残差审计_frozen.png"

d = Document(str(src))
psr = json.load(open(str(RD / "frozen_patient_specific_r_v2.json")))
summary = psr["summary"]
boot = psr["bootstrap_seed42"]
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# 1. 综合段更新（段106）
# ============================================================
new_summary = (
    "总体而言，6 h SOFA 轨迹以短时状态持续为主，69.2% 的锚点总 SOFA 未发生变化，"
    "使 persistence baseline 在全部锚点的总体 MAE 上优于两个学习模型；"
    "然而在后来实际发生 SOFA 状态转变的锚点中，Transformer 和 PLF-OGT 均明显优于 persistence，"
    "尽管 Transformer 的轨迹误差仍低于 PLF-OGT。"
    "固定模型通路对照进一步表明 PLF-OGT 功能性利用了锚点后的实际治疗信息，"
    "但该增益主要集中于心血管 SOFA 分量（80.1%）。"
    "Concept 与 residual pathways 均提供预测信息但并未完全分离；"
    "patient-specificity negative controls 表明 residual representation 携带不可由其他患者状态"
    "或群体平均替代的患者特异预测信息（matched 相对 shuffled 的 AUPRC 优势 +0.144，"
    "95% CI +0.119 至 +0.171），但单个 residual slot 不具有稳定可识别的临床语义。"
    "MIMIC-IV 独立再开发重现了完整 S+R representation 相对于通路阻断条件的优势，"
    "但未复制 GMUICU 中 residual-only 相对 concept-only 的排序。"
    "因此，PLF-OGT 的主要增量在于 structured auditability 和 patient-specific predictive representation，"
    "而非相对于标准 Transformer 的预测精度优势。"
)
set_para(d.paragraphs[106], new_summary)
rev.append("1. 综合段更新(纳入patient-specificity)")

# ============================================================
# 2. Table 3 (表2) 补四种负控行
# ============================================================
t3 = None
for ti, t in enumerate(d.tables):
    hdr = t.rows[0].cells[0].text if t.rows else ""
    if "审计项" in hdr:
        t3 = t
        break

if t3:
    # 追加行: patient-specificity 四条件
    conditions_info = [
        ("R patient-specificity: matched", f"{summary['matched']['auprc_mean']:.3f}", "正确R（基线）"),
        ("R patient-specificity: shuffled", f"{summary['shuffled']['auprc_mean']:.3f} (Δ−{abs(summary['shuffled']['delta_auprc']):.3f})", "换患者R，CI不跨零"),
        ("R patient-specificity: mean", f"{summary['mean']['auprc_mean']:.3f} (Δ−{abs(summary['mean']['delta_auprc']):.3f})", "群体平均R"),
        ("R patient-specificity: query-only", f"{summary['query_only']['auprc_mean']:.3f} (Δ−{abs(summary['query_only']['delta_auprc']):.3f})", "清零R内容"),
    ]
    for audit_item, result, note in conditions_info:
        row = t3.add_row()
        row.cells[0].paragraphs[0].add_run(audit_item)
        row.cells[1].paragraphs[0].add_run(result)
        row.cells[2].paragraphs[0].add_run(note)
    rev.append("2. Table 3 补4行patient-specificity")

# ============================================================
# 3. S6 补充表 - 追加 patient-specificity 行
# ============================================================
for ti, t in enumerate(d.tables):
    hdr = " ".join(c.text for c in t.rows[0].cells) if t.rows else ""
    if "遮蔽" in hdr or "槽匹配" in hdr:
        # 追加行
        for cond, label in [("matched", "R 替换：matched（正确R）"),
                            ("shuffled", "R 替换：shuffled（换患者R）"),
                            ("mean", "R 替换：population mean"),
                            ("query_only", "R 替换：query-only（清零R）")]:
            s = summary[cond]
            ci_info = boot.get(cond, {})
            ci_str = f"CI {ci_info['delta_auprc_ci'][0]:+.3f}至{ci_info['delta_auprc_ci'][1]:+.3f}" if ci_info else "—"
            row = t.add_row()
            row.cells[0].paragraphs[0].add_run(label)
            row.cells[1].paragraphs[0].add_run(f"AUPRC {s['auprc_mean']:.3f}±{s['auprc_std']:.3f}")
            if len(row.cells) >= 2:
                notes = f"macroMAE {s['macro_mae_mean']:.3f}"
                if cond != "matched" and ci_info:
                    notes += f"；matched−{cond} ΔAUPRC {ci_info['delta_auprc_point']:+.3f} ({ci_str})"
                row.cells[-1].paragraphs[0].add_run(notes)
        rev.append(f"3. S6(表{ti}) 补4行patient-specificity")
        break

# ============================================================
# 4. Figure 4 说明更新
# ============================================================
rep_runs(d.paragraphs[115],
    "代理恢复与随机映射 + S/R 2×2 + S/R 重叠 + 成功槽与失败槽。",
    "代理恢复与随机映射 + S/R 2×2 + patient-specificity negative controls + 成功槽与失败槽。")
rep_runs(d.paragraphs[174],
    "代理恢复与随机映射 + S/R 2×2 + S/R 重叠 + 成功槽与失败槽。",
    "代理恢复与随机映射 + S/R 2×2 + patient-specificity negative controls + 成功槽与失败槽。")
rev.append("4. Figure 4 说明更新")

d.save(str(out))

# ============================================================
# 5. 替换 Figure 4 图片
# ============================================================
z = zipfile.ZipFile(str(out), "r")
data = {n: z.read(n) for n in z.namelist()}
z.close()

# Figure 4 对应的 image — 需要找到它
# docx 里 Figure 4 可能还没有图片（之前是占位符）
# 检查有没有 image3.png 或更多
media_files = [n for n in data if "media" in n]
print(f"media files: {media_files}")

# 如果有 image3.png 或更多, 替换最后一个
# 如果只有 image1/2, Figure 4 还没有图片, 需要插入
# 检查 shape 数量
d2 = Document(str(out))
print(f"InlineShapes: {len(d2.inline_shapes)}")

# 如果只有2个shape (Fig2+Fig3), Fig4还是占位符
# 不能简单替换, 需要用 python-docx 插入图片到段175
# 但段175是 "〔图 4 待绘制后插入。〕"
p175 = d2.paragraphs[175]
if "待绘制" in p175.text:
    # 清空占位符, 插入图片
    for run in p175.runs:
        run.text = ""
    from docx.shared import Inches
    run = p175.add_run()
    run.add_picture(str(FIG), width=Inches(6.5))
    d2.save(str(out))
    rev.append("5. Figure 4 图片插入(原占位符替换)")
else:
    # 尝试替换已有图片
    for n in data:
        if "image3" in n or "image4" in n:
            data[n] = open(str(FIG), "rb").read()
            rev.append(f"5. Figure 4 图片替换: {n}")
            break
    with zipfile.ZipFile(str(out), "w", zipfile.ZIP_DEFLATED) as zo:
        for n in data:
            zo.writestr(n, data[n])

print(f"\n修订完成:")
for r in rev: print(f"  {r}")
print(f"保存: {out}")
