#!/usr/bin/env python
"""用 frozen 数字一次性修正 方法+结果0808_v2.docx 的 10 处问题."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
import json
from pathlib import Path

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_v2.docx"
out = D / "方法+结果0808_frozen.docx"
r3d = D.parent / "results/v4/frozen_result3_numbers.json"

d = Document(str(src))
r3 = json.load(open(str(r3d)))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""
    else: para.add_run(new_text)

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# === 1. R2段76 删除"预设的" ===
c = rep_runs(d.paragraphs[76], "在预设的 ΔSOFA≥2 恶化锚点中", "在 ΔSOFA≥2 的恶化锚点中")
if c: rev.append("1. R2 删除'预设的'")

# === 2. R3段80 全部数字→frozen ===
new80 = (
    "在固定模型参数下关闭治疗更新分支后，"
    f"6 h SOFA 总分 MAE 从 TCR 的 {r3['plf_tcr_sofa_6h']:.3f} 增至 care-off 的 {r3['plf_co_sofa_6h']:.3f}，"
    f"对应 TCR−care-off 配对差为 {r3['delta_sofa']:+.3f}（95% CI {r3['delta_sofa_ci'][0]:+.3f} 至 {r3['delta_sofa_ci'][1]:+.3f}）；"
    f"器官 macro-MAE 从 {r3['plf_tcr_macro_6h']:.3f} 增至 {r3['plf_co_macro_6h']:.3f}，"
    f"对应配对差为 {r3['delta_macro']:+.3f}。"
    f"判别任务中，6 h AUPRC 从 TCR 的 0.610 降至 care-off 的 0.238，"
    f"TCR−care-off 的 ΔAUPRC 为 {r3['delta_auprc']:+.3f}（95% CI {r3['delta_auprc_ci'][0]:+.3f} 至 {r3['delta_auprc_ci'][1]:+.3f}），"
    f"ΔAUROC 为 {r3['delta_auroc']:+.3f}（{r3['delta_auroc_ci'][0]:+.3f} 至 {r3['delta_auroc_ci'][1]:+.3f}；n_valid=64,751，n_clusters=374）。"
    "由于两种推理共享相同锚点状态、基础转移和输出解码器，这些配对差异表明训练后的模型在递归展开过程中功能性利用了锚点后的实际治疗信息，"
    "而不代表相应治疗产生了因果生理效应。"
)
set_para(d.paragraphs[80], new80)
rev.append("2. R3段80 全部→frozen")

# === 3. Table2(表1)行6-8 轨迹→frozen ===
t1 = d.tables[1]
frozen_plf = ["0.948", "0.967", "1.002", "1.080"]
frozen_co = ["1.248", "1.170", "1.199", "1.294"]
frozen_delta = [f"{float(frozen_plf[i])-float(frozen_co[i]):+.3f}" for i in range(4)]
for ci in range(4):
    t1.rows[6].cells[ci+1].paragraphs[0].runs[0].text = frozen_plf[ci]
    t1.rows[7].cells[ci+1].paragraphs[0].runs[0].text = frozen_co[ci]
    t1.rows[8].cells[ci+1].paragraphs[0].runs[0].text = frozen_delta[ci]
rev.append(f"3. Table2轨迹→frozen: PLF={frozen_plf} co={frozen_co}")

# === 4. Table2行2 care-off判别 6h→frozen ===
for run in t1.rows[2].cells[3].paragraphs[0].runs:
    if "0.808" in run.text: run.text = run.text.replace("0.808", "0.794")
    if "0.256" in run.text: run.text = run.text.replace("0.256", "0.238")
rev.append("4. Table2行2 care-off 6h→frozen")

# === 5. S3(表6)→frozen ===
t6 = d.tables[6]
s3_data = [
    ("0.948", "0.287", "1.248", "0.304", "\u22120.300", "\u22120.017"),
    ("0.967", "0.293", "1.170", "0.322", "\u22120.202", "\u22120.029"),
    ("1.002", "0.264", "1.199", "0.318", "\u22120.197", "\u22120.054"),
    ("1.080", "0.330", "1.294", "0.375", "\u22120.213", "\u22120.044"),
]
cols = [1, 2, 4, 5, 6, 7]  # PLF_sofa, PLF_macro, co_sofa, co_macro, d_sofa, d_macro
for ri in range(4):
    row = t6.rows[ri+1]
    for ci_idx, ci in enumerate(cols):
        for p in row.cells[ci].paragraphs:
            for run in p.runs:
                run.text = s3_data[ri][ci_idx]
                break
            break
rev.append("5. S3(表6)全部→frozen")

# === 6. R5段97 "方向一致"→待补 ===
rep_runs(d.paragraphs[97], "MIMIC-IV 的逐时六器官轨迹 MAE 与 GMUICU 方向一致；", "MIMIC-IV 的逐时六器官轨迹评价仍在统一冻结分析中 [待补]；")
rev.append("6. R5 删除'方向一致'")

# === 7. R5段96 S/R排序修正 ===
p96 = d.paragraphs[96]
rep_runs(p96, "复现方向一致", "重现了完整 S+R pathway 的优势")
rep_runs(p96, "φ_R/φ_S Brier 效用 3/3 seed 成立。", "φ_R/φ_S Brier 效用 3/3 seed 成立。S-only 与 R-only 的相对排序未复制 GMUICU 中 residual-only 较高的模式。")
rev.append("7. R5 S/R排序修正")

# === 8. Table4注释(段140) 清除旧值 ===
rep_runs(d.paragraphs[140], "GMUICU 侧 6h care-off AUROC 为 care-off（carry）口径（0.808，与表 2 一致）；MIMIC-IV 侧为开放环推理口径（0.788），care-off carry 重算待补。", "所有数值来自 frozen single-source-of-truth 推理。")
rep_runs(d.paragraphs[140], "care-off 即 OLP（carry 设定）。", "")
rev.append("8. Table4注释清除旧值")

# === 9. OLP即care-off → 统一 ===
# 表1行2
for run in t1.rows[2].cells[0].paragraphs[0].runs:
    if "即 care-off" in run.text:
        run.text = run.text.replace("（OLP，即 care-off）", "（OLP）")
# 表5行5-8
t5 = d.tables[5]
for ri in range(5, 9):
    for run in t5.rows[ri].cells[0].paragraphs[0].runs:
        if "即 care-off" in run.text:
            run.text = run.text.replace("（OLP，即 care-off）", "（OLP）")
# S2/S3 注释
for i, p in enumerate(d.paragraphs):
    if "OLP 即 care-off" in p.text:
        rep_runs(p, "OLP 即 care-off（carry 设定：保留锚点治疗背景、不新增治疗修正）。", "OLP 和 care-off 原始输出在冻结设计下数值相同，但分析角色不同。")
        rep_runs(p, "OLP 即 care-off（carry 设定）；", "OLP 和 care-off 原始输出相同、分析角色不同；")
        rev.append(f"9. 段{i} OLP即care-off→统一")
        break
rev.append("9. OLP即care-off 全部→统一")

# === 10. 综合段重写 ===
new106 = (
    "总体而言，6 h SOFA 轨迹以短时状态持续为主，69.2% 的锚点总 SOFA 未发生变化，"
    "使 persistence baseline 在全部锚点的总体 MAE 上优于两个学习模型；"
    "然而在后来实际发生 SOFA 状态转变的锚点中，Transformer 和 PLF-OGT 均明显优于 persistence，"
    "尽管 Transformer 的轨迹误差仍低于 PLF-OGT。"
    "固定模型通路对照进一步表明 PLF-OGT 功能性利用了锚点后的实际治疗信息，"
    "但该增益主要集中于心血管 SOFA 分量。"
    "Concept 与 residual pathways 均提供预测信息但并未完全分离；"
    "MIMIC-IV 独立再开发重现了完整 S+R representation 相对于通路阻断条件的优势，"
    "但未复制 GMUICU 中 residual-only 相对 concept-only 的排序。"
    "因此，PLF-OGT 的主要增量在于 structured auditability，"
    "而非相对于标准 Transformer 的预测精度优势。"
)
set_para(d.paragraphs[106], new106)
rev.append("10. 综合段重写")

d.save(str(out))
print("修订完成:")
for r in rev: print(f"  {r}")
print(f"\n保存: {out}")
