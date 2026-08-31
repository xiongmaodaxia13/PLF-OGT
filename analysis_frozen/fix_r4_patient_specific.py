#!/usr/bin/env python
"""写入 Result 4 patient-specific 负控正文 + 更新标题."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from pathlib import Path

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_v2.docx"
out = D / "方法+结果0808_frozen_r4.docx"

d = Document(str(src))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""
    else: para.add_run(new_text)

# ============================================================
# 1. Result 4 标题升级
# ============================================================
set_para(d.paragraphs[85],
    "Result 4. Concept 与 residual 通路提供互补信息，residual representation 携带患者特异预测内容但缺乏稳定槽级语义")
rev.append("1. R4标题升级")

# ============================================================
# 2. 段91 [待补] → 审稿人提供的四段式正文
# ============================================================
text_r4 = (
    "为检验 residual representation 是否携带患者对应的信息，而不仅反映共享的 residual slot 结构或群体平均状态，"
    "我们进行了 patient-specificity negative controls。"
    "对于每个测试 ICU stay，matched 条件使用该患者自身的 residual state；"
    "shuffled 条件采用 stay-level derangement，将 residual state 替换为另一住院 episode 的状态并确保不发生同 stay 重配；"
    "population-mean 条件使用测试人群平均 residual state；query-only 条件则清零 residual content。"
    "三次独立训练汇总中，matched 条件的 6 h AUPRC 和 macro-MAE 分别为 0.567 和 0.315。"
    "将 residual state 替换为其他患者状态后，AUPRC 降至 0.410，macro-MAE 增至 0.928，"
    "相对 matched 的三-seed 平均变化分别为 −0.157 和 +0.613。三个独立训练均显示相同方向的性能下降。"
    "\n\n"
    "在 seed 42 的 ICU-stay 级配对 cluster bootstrap（2,000 次）中，"
    "matched 相对 shuffled 的 AUPRC 优势为 +0.144（95% CI +0.119 至 +0.171），"
    "shuffled 相对 matched 的 macro-MAE 增量 95% CI 为 +0.474 至 +0.563；两项区间均未跨 0。"
    "Population-mean 和 query-only 条件同样显示稳定的性能损失，"
    "matched 的 AUPRC 优势分别为 +0.100（+0.079 至 +0.123）和 +0.160（+0.133 至 +0.187）。"
    "\n\n"
    "值得注意的是，错误患者的 residual state 在判别和轨迹回归中均产生了最明显的总体损害："
    "shuffled AUPRC 低于 query-only（0.410 vs 0.423），macro-MAE 也明显更高（0.928 vs 0.746）。"
    "这表明 residual representation 的作用不仅来自 residual pathway 的存在；"
    "错误的患者—状态匹配可引入具有方向性的错配信息。"
    "Population-mean replacement 的 AUPRC 为 0.466，高于 query-only 的 0.423，"
    "但其 macro-MAE 为 0.756，与 query-only 的 0.746 接近，"
    "因此该结果仅提示 residual representation 可能同时包含一定群体共享结构，"
    "而不支持将其严格分解为独立的 population-level 与 patient-specific components。"
    "\n\n"
    "综合而言，这些负控支持 residual representation 携带患者特异的预测信息（patient-specific predictive information），"
    "且这种信息不能由群体平均状态或其他患者的 residual state 完全替代。"
    "然而，这种患者特异性并不意味着单个 residual slot 具有稳定、唯一或可识别的临床语义；"
    "结合 S–R 双向信息重叠和有限的跨 seed slot identity stability，"
    "本研究仅将 R 解释为 patient-specific predictive latent representation，"
    "而非已识别的患者特异病理生理机制。"
)
set_para(d.paragraphs[91], text_r4)
rev.append("2. R4段91 [待补]→四段式正文")

d.save(str(out))
print("修订完成:")
for r in rev: print(f"  {r}")
print(f"保存: {out}")
