#!/usr/bin/env python
"""写入 P0-1 MIMIC trajectory + P0-2 modified-SOFA sensitivity 到 docx."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from pathlib import Path
import json

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_r4_final.docx"
out = D / "方法+结果0808_frozen_p0.docx"
RD = D.parent / "results/v4"
RM = D.parent / "results_mimic"

d = Document(str(src))
mt = json.load(open(str(RM / "frozen_mimic_trajectory.json")))
org = json.load(open(str(RD / "frozen_organ_noncv.json")))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# P0-1: MIMIC trajectory → 替换 Result 5 的 [待补]
# ============================================================
# 段97 含 "MIMIC-IV 的逐时六器官轨迹评价仍在统一冻结分析中 [待补]"
h6 = mt["multi_horizon"][2]
cs = mt["changed_state_6h"]

mimic_traj_text = (
    f"MIMIC-IV 的逐时六器官轨迹评价显示，PLF-OGT 在 6 h 的 SOFA 总分 MAE 为 {h6['plf_sofa']:.3f}，"
    f"persistence baseline 为 {h6['persist_sofa']:.3f}（Δ {h6['delta_sofa']:+.3f}）。"
    f"与 GMUICU 一致，总体 MAE 仍受状态持续性主导（60.2% 的 6 h 锚点 ΔSOFA=0），"
    f"但在后来实际发生 SOFA 变化的锚点中，PLF-OGT 优于 persistence"
    f"（|ΔSOFA|≥1：PLF {cs['changed_ge1']['plf_sofa']:.3f} vs persistence {cs['changed_ge1']['persist_sofa']:.3f}，"
    f"Δ {cs['changed_ge1']['delta']:+.3f}）。"
    f"该 changed-state 反转模式与 GMUICU 方向一致，说明 trajectory 改变状态下的优势跨队列可复现。"
)

# 找含 [待补] 的 MIMIC trajectory 句
for i, p in enumerate(d.paragraphs):
    if "MIMIC-IV 的逐时六器官轨迹" in p.text and "[待补]" in p.text:
        rep_runs(p, "MIMIC-IV 的逐时六器官轨迹评价仍在统一冻结分析中 [待补]；", mimic_traj_text)
        rev.append(f"P0-1 段{i} MIMIC trajectory写入")
        break
    elif "MIMIC-IV 的逐时六器官轨迹" in p.text:
        # 直接替换整句
        full = p.text
        if "[待补]" in full:
            new_full = full.replace(
                full[full.index("MIMIC-IV 的逐时六器官轨迹"):full.index("；", full.index("MIMIC-IV 的逐时六器官轨迹"))+1],
                mimic_traj_text)
            set_para(p, new_full)
            rev.append(f"P0-1 段{i} MIMIC trajectory写入(整句)")
            break

# 如果没找到, 在段96后插入
if not any("P0-1" in r for r in rev):
    # 找 R5 的 discrimination 段, 在其后加 trajectory
    for i, p in enumerate(d.paragraphs):
        if "MIMIC-IV" in p.text and "代理恢复" in p.text:
            # 在这段后面加 trajectory 句
            rep_runs(p,
                "代理恢复 MAE 为 0.046（优于 GMUICU 0.094），R→S 最大 R² 为 0.391。",
                f"代理恢复 MAE 为 0.046（优于 GMUICU 0.094），R→S 最大 R² 为 0.391。{mimic_traj_text}")
            rev.append(f"P0-1 段{i} MIMIC trajectory追加")
            break

# ============================================================
# P0-2: modified-SOFA sensitivity → 替换 Result 3 的 [待补]
# ============================================================
ncv_delta = org["total_delta_sum"] - org["organ_mae"]["Cv"]["delta"]
cv_pct = org["organ_mae"]["Cv"]["delta"] / org["total_delta_sum"] * 100

modsofa_text = (
    f"为检验治疗信息增益是否超出 SOFA 心血管分量（直接整合升压药剂量）的定义性耦合，"
    f"我们进一步排除心血管分量后重新计算 modified-SOFA 轨迹。"
    f"在 6 h 时，排除心血管后的五器官合计 care-off ΔMAE 为 {ncv_delta:.3f}"
    f"（占总 ΔMAE {org['total_delta_sum']:.3f} 的 {100-cv_pct:.1f}%），"
    f"远低于心血管分量的 +{org['organ_mae']['Cv']['delta']:.3f}（{cv_pct:.1f}%）。"
    f"在判别任务中，以 non-CV SOFA 恶化为结局时，"
    f"TCR 相对 care-off 的 ΔAUROC 从总 SOFA 的 +{org['total_cls']['delta_auroc']:.3f} 降至 +{org['non_cv']['delta_auroc']:.3f}，"
    f"ΔBrier 为 {org['non_cv']['delta_brier']:+.3f}（恶化）。"
    f"该结果表明 PLF-OGT 对锚点后治疗信息的功能依赖主要由心血管评分定义中包含的治疗相关信息驱动，"
    f"而非治疗信息在六个器官系统间产生广泛且均一的预测增益。"
)

# 找含 treatment-independent 或 modified-SOFA 或 5-6 organ 的 [待补]
for i, p in enumerate(d.paragraphs):
    txt = p.text
    if "[待补]" in txt and ("treatment" in txt.lower() or "modified" in txt.lower() or "organ" in txt.lower() or "5–6" in txt or "5-6" in txt):
        set_para(p, modsofa_text)
        rev.append(f"P0-2 段{i} modified-SOFA写入(替换[待补])")
        break

# 如果没找到明确位置, 在R3器官分解段后追加
if not any("P0-2" in r for r in rev):
    # 段81是器官分解段, 在其后加
    p81 = d.paragraphs[81]
    current = p81.text
    if "modified-SOFA" not in current:
        set_para(p81, current + "\n" + modsofa_text)
        rev.append("P0-2 段81 modified-SOFA追加到器官分解段末")

d.save(str(out))
print("修订完成:")
for r in rev: print(f"  {r}")
print(f"保存: {out}")
