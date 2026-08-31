#!/usr/bin/env python
"""替换 Figure 3 为 frozen 版 + 修 S4 CI列 + R3措辞 + typo."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from pathlib import Path
import zipfile, shutil

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen_final.docx"
out = D / "方法+结果0808_frozen_v2.docx"
FIG = D.parent / "figures/图3_器官级范围_frozen.png"

d = Document(str(src))
rev = []

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# 1. 替换 Figure 3 图片 (inline shape)
#    docx 有 2 个 inline shape, Figure 2=shape0, Figure 3=shape1
# ============================================================
# python-docx 不直接支持替换图片, 用 zip 操作替换 word/media/
# 先找到 docx 里 Figure 3 对应的 media 文件

z = zipfile.ZipFile(str(src), "r")
names = z.namelist()
media_files = [n for n in names if n.startswith("word/media/")]
z.close()
print(f"docx 内 media 文件: {media_files}")

# 方法: 直接在 docx 的 zip 里替换 image 文件
# 需要知道哪个 image 是 Figure 3
# shape1 (第二个图片) = Figure 3
# 通常 image1.png=Figure2, image2.png=Figure3 (按插入顺序)

# 简单做法: 用 python-docx 的 inline_shape 替换图片数据
# 或者直接用 zip 操作

# 先用 python-docx 看 shape 关系
for si, s in enumerate(d.inline_shapes):
    rId = s._inline.graphic.graphicData.pic.blipFill.blip.embed
    print(f"  shape{si}: rId={rId}")

# 获取关系映射
rels = d.part.rels
for si, s in enumerate(d.inline_shapes):
    rId = s._inline.graphic.graphicData.pic.blipFill.blip.embed
    if rId in rels:
        target = rels[rId].target_ref
        print(f"  shape{si} -> {target}")

rev.append("Figure 3 图片待替换(见下方 zip 操作)")

# ============================================================
# 2. S4 no-residual CI 列位置修正
# ============================================================
# 找 S4 表格 (含 "no_residual" 或 "移除残差" 的表)
for ti, t in enumerate(d.tables):
    for ri, row in enumerate(t.rows):
        txt = " ".join(c.text for c in row.cells)
        if "移除残差" in txt and ("3" in txt and "独立" in txt):
            # 这行: AUROC / AUPRC / ΔAUROC / ΔAUPRC
            # 当前: AUPRC=0.543, ΔAUPRC=-0.026(95% CI 0.514-0.571) ← CI放错
            # 正确: AUPRC=0.543(95% CI 0.514-0.571), ΔAUPRC=-0.026
            cells = row.cells
            # 找 AUPRC 列 (通常第3列 index=2)
            for ci in range(len(cells)):
                cell_text = cells[ci].text.strip()
                if "0.543" in cell_text and "0.514" not in cell_text:
                    # AUPRC 列: 改为 0.543 (0.514-0.571)
                    for p in cells[ci].paragraphs:
                        for run in p.runs:
                            if "0.543" in run.text:
                                run.text = "0.543（0.514–0.571）"
                                break
                        break
                    break
                elif "0.543" in cell_text and "0.514" in cell_text:
                    # CI 已经在 AUPRC 列了, 检查 ΔAUPRC 列
                    pass
            # 找 ΔAUPRC 列: 把 (0.514-0.571) 从 ΔAUPRC 列移走
            for ci in range(len(cells)):
                cell_text = cells[ci].text.strip()
                if "0.514" in cell_text and ("-0.026" in cell_text or "−0.026" in cell_text):
                    # ΔAUPRC 列: 只保留 -0.026
                    for p in cells[ci].paragraphs:
                        for run in p.runs:
                            if "0.514" in run.text:
                                run.text = run.text.replace("（95% CI 0.514–0.571）", "").replace("(0.514-0.571)", "")
                                run.text = run.text.replace("0.514–0.571", "").strip()
                                break
                        break
            rev.append(f"2. S4 表{ti}行{ri} CI列位置修正")
            break

# ============================================================
# 3. R3 段81 ΔAUROC 措辞: "care-off的ΔAUROC" → "TCR相对care-off的ΔAUROC"
# ============================================================
p81 = d.paragraphs[81]
rep_runs(p81, "care-off 的 ΔAUROC 为 +0.081", "TCR 相对 care-off 的 ΔAUROC 为 +0.081")
rep_runs(p81, "care-off 的 ΔAUROC", "TCR 相对 care-off 的 ΔAUROC")
rev.append("3. R3段81 ΔAUROC措辞修正")

# ============================================================
# 4. GMUICU typo: "GMICU" → "GMUICU"
# ============================================================
for i, p in enumerate(d.paragraphs):
    c = rep_runs(p, "GMICU", "GMUICU")
    if c: rev.append(f"4. 段{i} GMICU→GMUICU")

# ============================================================
# 5. "机制估计量" → "功能通路比较指标"
# ============================================================
for i, p in enumerate(d.paragraphs):
    c1 = rep_runs(p, "机制估计量", "功能通路比较指标")
    if c1: rev.append(f"5. 段{i} 机制估计量→功能通路比较指标")

d.save(str(out))

# ============================================================
# 6. 替换 Figure 3 图片 (zip 操作)
# ============================================================
# 找 Figure 3 对应的 image part
# 从 shape1 的 rId 找 target
fig3_target = None
for si, s in enumerate(d.inline_shapes):
    if si == 1:  # Figure 3
        rId = s._inline.graphic.graphicData.pic.blipFill.blip.embed
        if rId in d.part.rels:
            fig3_target = d.part.rels[rId].target_ref
            break

if fig3_target:
    print(f"\nFigure 3 对应: word/{fig3_target}")
    # 用 zip 替换
    import zipfile
    z_out = zipfile.ZipFile(str(out), "r")
    data = {n: z_out.read(n) for n in z_out.namelist()}
    z_out.close()

    media_path = f"word/{fig3_target}" if not fig3_target.startswith("word/") else fig3_target
    if media_path in data:
        data[media_path] = open(str(FIG), "rb").read()
        with zipfile.ZipFile(str(out), "w", zipfile.ZIP_DEFLATED) as zo:
            for n in data:
                zo.writestr(n, data[n])
        rev.append(f"6. Figure 3 图片替换: {media_path}")
    else:
        rev.append(f"6. Figure 3 media 未找到: {media_path}")
else:
    rev.append("6. Figure 3 rId 未找到, 尝试替换 word/media/image2.png")
    # fallback: 直接替换 image2.png
    z_out = zipfile.ZipFile(str(out), "r")
    data = {n: z_out.read(n) for n in z_out.namelist()}
    z_out.close()
    for n in data:
        if "image2" in n and "media" in n:
            data[n] = open(str(FIG), "rb").read()
            rev.append(f"6. 替换 {n}")
            break
    with zipfile.ZipFile(str(out), "w", zipfile.ZIP_DEFLATED) as zo:
        for n in data:
            zo.writestr(n, data[n])

print(f"\n修订完成:")
for r in rev: print(f"  {r}")
print(f"保存: {out}")
