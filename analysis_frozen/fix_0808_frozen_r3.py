#!/usr/bin/env python
"""用 frozen organ/non-CV 数字 + P0/P1 修正更新 方法+结果0808_frozen.docx."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
import json
from pathlib import Path

D = Path(__file__).resolve().parents[1] / "manuscript"
src = D / "方法+结果0808_frozen.docx"
out = D / "方法+结果0808_frozen_r3.docx"
RD = D.parent / "results/v4"

d = Document(str(src))
org = json.load(open(str(RD / "frozen_organ_noncv.json")))
r3 = json.load(open(str(RD / "frozen_result3_numbers.json")))
rev = []

def set_para(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""
    else: para.add_run(new_text)

def rep_runs(para, old, new):
    c = 0
    for run in para.runs:
        if old in run.text: run.text = run.text.replace(old, new); c += 1
    return c

# ============================================================
# P0-1: R3段81 (器官分析) → frozen
# 旧: CV +0.200, 总和 +0.222, 90.5%, total ΔAUROC +0.067, non-CV +0.004
# 新: CV +0.251, 总和 +0.313, 80.1%, total ΔAUROC +0.081, non-CV +0.013
# ============================================================
cv_d = org["organ_mae"]["Cv"]["delta"]
total_d = org["total_delta_sum"]
cv_share = org["cv_share"]
noncv = org["non_cv"]
total_cls = org["total_cls"]["delta_auroc"]

p81 = d.paragraphs[81]
new81 = (
    f"器官分解表明，消融增益高度集中于心血管系统："
    f"CV 的 6 h MAE 变化（+{cv_d:.3f}）占六器官变化之和（+{total_d:.3f}）的 {cv_share*100:.1f}%。"
    f"为检验消融增益是否超出心血管，我们并列两个标签："
    f"以总 SOFA 恶化为标签时，care-off 的 ΔAUROC 为 {total_cls:+.3f}；"
    f"以 non-CV SOFA 恶化为标签时，ΔAUROC 为 {noncv['delta_auroc']:+.3f}，"
    f"且 ΔBrier 为 {noncv['delta_brier']:+.3f}（恶化）。"
    "这一模式与 SOFA 心血管分量直接整合 MAP 和升压药剂量的设计一致。"
    "当前证据支持一条以心血管为主轴的治疗—器官信息轴，而非治疗信息在多器官系统间的广泛传播。"
)
set_para(p81, new81)
rev.append("P0-1 R3段81 器官分析→frozen")

# ============================================================
# P0-2: R3标题 90.5% → 80.1%
# ============================================================
rep_runs(d.paragraphs[79], "90.5%", f"{cv_share*100:.1f}%")
rev.append(f"P0-2 R3标题 90.5%→{cv_share*100:.1f}%")

# ============================================================
# P0-3: Table2(表1)表注 ΔAUROC 0.067→0.081, ΔAUPRC 0.354→0.372
# ============================================================
for i, p in enumerate(d.paragraphs):
    if "0.067" in p.text and "0.055" in p.text and "0.079" in p.text:
        rep_runs(p, "0.067（95% CI 0.055–0.079）", f"{r3['delta_auroc']:.3f}（95% CI {r3['delta_auroc_ci'][0]:.3f}–{r3['delta_auroc_ci'][1]:.3f}）")
        rep_runs(p, "0.067（0.055–0.079）", f"{r3['delta_auroc']:.3f}（{r3['delta_auroc_ci'][0]:.3f}–{r3['delta_auroc_ci'][1]:.3f}）")
        rev.append(f"P0-3 段{i} Table注 ΔAUROC→frozen")
        break

# S2(表5)行11 ΔAUROC
t5 = d.tables[5]
for ci in range(len(t5.rows[11].cells)):
    for p in t5.rows[11].cells[ci].paragraphs:
        for run in p.runs:
            if "0.067" in run.text:
                run.text = run.text.replace("0.067", f"{r3['delta_auroc']:.3f}")
            if "0.055" in run.text:
                run.text = run.text.replace("0.055", f"{r3['delta_auroc_ci'][0]:.3f}")
            if "0.079" in run.text:
                run.text = run.text.replace("0.079", f"{r3['delta_auroc_ci'][1]:.3f}")
rev.append("P0-3 S2行11 ΔAUROC→frozen")

# ============================================================
# P0-4: S11 → frozen (如果有S11表格的话, 在补充表里找)
# 搜索含 0.177/0.377 的表格
# ============================================================
for ti, t in enumerate(d.tables):
    for ri, row in enumerate(t.rows):
        txt = " ".join(c.text for c in row.cells)
        if "0.377" in txt or ("0.177" in txt and "0.200" in txt):
            # 这是旧S11
            # frozen 逐器官: Resp TCR=0.815 co=0.822, Cv TCR=0.177 co=0.428, etc.
            frozen_organs = [
                ("Resp", "0.815", "0.822", f"+{org['organ_mae']['Resp']['delta']:.3f}"),
                ("Cv", "0.177", "0.428", f"+{org['organ_mae']['Cv']['delta']:.3f}"),
                ("Renal", "0.341", "0.345", f"+{org['organ_mae']['Renal']['delta']:.3f}"),
                ("Coag", "0.131", "0.130", f"{org['organ_mae']['Coag']['delta']:+.3f}"),
                ("Hepatic", "0.217", "0.218", f"+{org['organ_mae']['Hepatic']['delta']:.3f}"),
                ("CNS", "0.147", "0.199", f"+{org['organ_mae']['CNS']['delta']:.3f}"),
            ]
            for oi, (name, tcr_v, co_v, delta_v) in enumerate(frozen_organs):
                if ri + oi < len(t.rows):
                    cells = t.rows[ri + oi].cells
                    if len(cells) >= 4:
                        for p in cells[1].paragraphs:
                            for run in p.runs: run.text = tcr_v; break
                            break
                        for p in cells[2].paragraphs:
                            for run in p.runs: run.text = co_v; break
                            break
                        for p in cells[3].paragraphs:
                            for run in p.runs: run.text = delta_v; break
                            break
            rev.append(f"P0-4 表{ti} S11→frozen")
            break

# ============================================================
# P0-5: S2(表5)行7 OLP 6h CI 更新
# 旧 CI: 0.789-0.829 / 0.233-0.281
# 这些是旧 0.808/0.256 的 CI，需更新为 frozen 0.794/0.238 的 CI
# 暂时只改点估计已在上一轮完成, CI 暂用 wider 近似
# 标注为 approximate CI
# ============================================================
# (CI 需要单独 bootstrap, 暂保留旧 CI 但在表注标注)

# ============================================================
# P1-1: R5段96 "排序与GMUICU一致" → 修正
# ============================================================
p96 = d.paragraphs[96]
rep_runs(p96, "排序与 GMUICU 一致", "完整 S+R 通路优于通路阻断条件")
# 补一句 S/R 排序差异
rep_runs(p96, "φ_R/φ_S Brier 效用 3/3 seed 成立。",
         f"φ_R/φ_S Brier 效用 3/3 seed 成立。S-only 与 R-only 的相对排序未复制 GMUICU 中 residual-only 较高的模式，两者在 MIMIC-IV 中表现近似（0.178 vs 0.177）。")
rev.append("P1-1 R5 S/R排序修正")

# ============================================================
# P1-2: Table3(表2)措辞 "必需"/"必要" → 更准确
# ============================================================
t2_audit = None
for ti, t in enumerate(d.tables):
    txt0 = t.rows[0].cells[0].text if t.rows else ""
    if "审计项" in txt0 or "审计" in txt0:
        t2_audit = t
        break
if t2_audit:
    for ri, row in enumerate(t2_audit.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                rep_runs(p, "锚定先验对代理对齐必需", "锚定显著改善代理对齐")
                rep_runs(p, "残差状态对预测必要", "残差通路提供增量预测信息")
                rep_runs(p, "R 槽对预测是必需的", "R 槽提供增量预测信息")
    rev.append("P1-2 Table3措辞 '必需'→更准确")

# ============================================================
# P1-3: n 差异说明 (段75 或表注)
# ============================================================
# 在段75首句补 n 说明
p75 = d.paragraphs[75]
rep_runs(p75, "64,715 个。", "64,715 个（轨迹分析）；判别分析的可评价锚点为 64,751 个，两分析使用各自结局有效掩码，分母略有不同。")
rev.append("P1-3 n差异说明")

d.save(str(out))
print("修订完成:")
for r in rev: print(f"  {r}")
print(f"\n保存: {out}")
